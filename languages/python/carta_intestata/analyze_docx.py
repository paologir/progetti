#!/usr/bin/env python3
"""
Script per analizzare il file Modello 2025.docx ed estrarre informazioni su:
- Intestazione (header)
- Piè di pagina (footer)
- Stili utilizzati
- Margini e formattazione
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
import os
import zipfile
import xml.etree.ElementTree as ET

def safe_get_margin(section, margin_name):
    """Ottiene un margine in modo sicuro, gestendo errori di conversione"""
    try:
        margin_value = getattr(section, margin_name)
        if margin_value is not None:
            return margin_value / 914400  # Conversione da EMU a cm
    except (ValueError, AttributeError):
        pass
    return None

def analyze_docx(file_path):
    """Analizza un file DOCX e restituisce informazioni dettagliate"""
    
    if not os.path.exists(file_path):
        print(f"Errore: Il file {file_path} non esiste")
        return
    
    try:
        doc = Document(file_path)
        
        print("=" * 80)
        print(f"ANALISI DEL FILE: {file_path}")
        print("=" * 80)
        
        # 1. ANALISI SEZIONI E MARGINI
        print("\n1. SEZIONI E MARGINI DEL DOCUMENTO:")
        print("-" * 40)
        
        for i, section in enumerate(doc.sections):
            print(f"\nSezione {i + 1}:")
            
            # Usa la funzione safe per ottenere i margini
            top_margin = safe_get_margin(section, 'top_margin')
            if top_margin is not None:
                print(f"  - Margine superiore: {top_margin:.2f} cm")
            
            bottom_margin = safe_get_margin(section, 'bottom_margin')
            if bottom_margin is not None:
                print(f"  - Margine inferiore: {bottom_margin:.2f} cm")
            
            left_margin = safe_get_margin(section, 'left_margin')
            if left_margin is not None:
                print(f"  - Margine sinistro: {left_margin:.2f} cm")
            
            right_margin = safe_get_margin(section, 'right_margin')
            if right_margin is not None:
                print(f"  - Margine destro: {right_margin:.2f} cm")
            
            header_distance = safe_get_margin(section, 'header_distance')
            if header_distance is not None:
                print(f"  - Margine intestazione: {header_distance:.2f} cm")
            
            footer_distance = safe_get_margin(section, 'footer_distance')
            if footer_distance is not None:
                print(f"  - Margine piè di pagina: {footer_distance:.2f} cm")
            
            page_width = safe_get_margin(section, 'page_width')
            if page_width is not None:
                print(f"  - Larghezza pagina: {page_width:.2f} cm")
            
            page_height = safe_get_margin(section, 'page_height')
            if page_height is not None:
                print(f"  - Altezza pagina: {page_height:.2f} cm")
            
            try:
                print(f"  - Orientamento: {'Verticale' if section.orientation == 0 else 'Orizzontale'}")
            except:
                pass
            
            # 2. ANALISI INTESTAZIONE (HEADER)
            print(f"\n2. INTESTAZIONE (Header) - Sezione {i + 1}:")
            print("-" * 40)
            
            header = section.header
            if header.paragraphs:
                for j, para in enumerate(header.paragraphs):
                    if para.text.strip():
                        print(f"  Paragrafo {j + 1}:")
                        print(f"    - Testo: '{para.text}'")
                        print(f"    - Allineamento: {para.alignment}")
                        if para.runs:
                            for k, run in enumerate(para.runs):
                                if run.text.strip():
                                    print(f"    - Run {k + 1}:")
                                    print(f"      - Testo: '{run.text}'")
                                    if run.font.name:
                                        print(f"      - Font: {run.font.name}")
                                    if run.font.size:
                                        print(f"      - Dimensione: {run.font.size.pt} pt")
                                    print(f"      - Grassetto: {run.font.bold}")
                                    print(f"      - Corsivo: {run.font.italic}")
                                    print(f"      - Sottolineato: {run.font.underline}")
                                    if run.font.color.rgb:
                                        print(f"      - Colore: RGB({run.font.color.rgb})")
                
                # Controlla se ci sono tabelle nell'intestazione
                if header.tables:
                    print(f"\n  Tabelle nell'intestazione: {len(header.tables)}")
                    for t, table in enumerate(header.tables):
                        print(f"    - Tabella {t + 1}: {len(table.rows)} righe x {len(table.columns)} colonne")
                        for r, row in enumerate(table.rows):
                            for c, cell in enumerate(row.cells):
                                if cell.text.strip():
                                    print(f"      - Cella [{r},{c}]: '{cell.text}'")
            else:
                print("  Nessuna intestazione definita in questa sezione")
            
            # 3. ANALISI PIÈ DI PAGINA (FOOTER)
            print(f"\n3. PIÈ DI PAGINA (Footer) - Sezione {i + 1}:")
            print("-" * 40)
            
            footer = section.footer
            if footer.paragraphs:
                for j, para in enumerate(footer.paragraphs):
                    if para.text.strip():
                        print(f"  Paragrafo {j + 1}:")
                        print(f"    - Testo: '{para.text}'")
                        print(f"    - Allineamento: {para.alignment}")
                        if para.runs:
                            for k, run in enumerate(para.runs):
                                if run.text.strip():
                                    print(f"    - Run {k + 1}:")
                                    print(f"      - Testo: '{run.text}'")
                                    if run.font.name:
                                        print(f"      - Font: {run.font.name}")
                                    if run.font.size:
                                        print(f"      - Dimensione: {run.font.size.pt} pt")
                                    print(f"      - Grassetto: {run.font.bold}")
                                    print(f"      - Corsivo: {run.font.italic}")
                                    print(f"      - Sottolineato: {run.font.underline}")
                                    if run.font.color.rgb:
                                        print(f"      - Colore: RGB({run.font.color.rgb})")
                
                # Controlla se ci sono tabelle nel piè di pagina
                if footer.tables:
                    print(f"\n  Tabelle nel piè di pagina: {len(footer.tables)}")
                    for t, table in enumerate(footer.tables):
                        print(f"    - Tabella {t + 1}: {len(table.rows)} righe x {len(table.columns)} colonne")
                        for r, row in enumerate(table.rows):
                            for c, cell in enumerate(row.cells):
                                if cell.text.strip():
                                    print(f"      - Cella [{r},{c}]: '{cell.text}'")
            else:
                print("  Nessun piè di pagina definito in questa sezione")
        
        # 4. ANALISI STILI DEL DOCUMENTO
        print("\n\n4. STILI DEFINITI NEL DOCUMENTO:")
        print("-" * 40)
        
        styles_used = set()
        for para in doc.paragraphs:
            if para.style.name:
                styles_used.add(para.style.name)
        
        print(f"Stili utilizzati nei paragrafi: {len(styles_used)}")
        for style_name in sorted(styles_used):
            print(f"  - {style_name}")
            # Prova a ottenere informazioni sullo stile
            try:
                style = doc.styles[style_name]
                if hasattr(style, 'font'):
                    if style.font.name:
                        print(f"    Font: {style.font.name}")
                    if style.font.size:
                        print(f"    Dimensione: {style.font.size.pt} pt")
            except:
                pass
        
        # 5. ANALISI CONTENUTO PRINCIPALE
        print("\n\n5. CONTENUTO PRINCIPALE DEL DOCUMENTO:")
        print("-" * 40)
        
        print(f"Numero totale di paragrafi: {len(doc.paragraphs)}")
        print(f"Numero totale di tabelle: {len(doc.tables)}")
        
        # Mostra i primi paragrafi non vuoti
        print("\nPrimi paragrafi con contenuto:")
        count = 0
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() and count < 5:
                count += 1
                print(f"\nParagrafo {i + 1}:")
                print(f"  - Testo: '{para.text[:100]}{'...' if len(para.text) > 100 else ''}'")
                print(f"  - Stile: {para.style.name}")
                print(f"  - Allineamento: {para.alignment}")
                
                # Analizza il primo run per informazioni sul font
                if para.runs:
                    run = para.runs[0]
                    if run.font.name:
                        print(f"  - Font: {run.font.name}")
                    if run.font.size:
                        print(f"  - Dimensione font: {run.font.size.pt} pt")
        
        # 6. PROPRIETÀ GENERALI DEL DOCUMENTO
        print("\n\n6. PROPRIETÀ DEL DOCUMENTO:")
        print("-" * 40)
        
        core_props = doc.core_properties
        if core_props.author:
            print(f"  - Autore: {core_props.author}")
        if core_props.title:
            print(f"  - Titolo: {core_props.title}")
        if core_props.subject:
            print(f"  - Oggetto: {core_props.subject}")
        if core_props.keywords:
            print(f"  - Parole chiave: {core_props.keywords}")
        if core_props.created:
            print(f"  - Data creazione: {core_props.created}")
        if core_props.modified:
            print(f"  - Ultima modifica: {core_props.modified}")
        
        print("\n" + "=" * 80)
        print("ANALISI COMPLETATA")
        print("=" * 80)
        
    except Exception as e:
        print(f"Errore durante l'analisi: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    # Percorso del file da analizzare
    file_path = "/opt/progetti/python/carta_intestata/Modello 2025.docx"
    analyze_docx(file_path)