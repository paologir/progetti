#!/usr/bin/env python3
"""
Script avanzato per analizzare in dettaglio il file Modello 2025.docx
"""

from docx import Document
from docx.shared import RGBColor
import zipfile
import xml.etree.ElementTree as ET
import os

def analyze_xml_content(file_path):
    """Analizza il contenuto XML del documento per informazioni aggiuntive"""
    print("\n7. ANALISI XML DETTAGLIATA:")
    print("-" * 40)
    
    try:
        with zipfile.ZipFile(file_path, 'r') as docx_zip:
            # Analizza styles.xml per informazioni sugli stili
            if 'word/styles.xml' in docx_zip.namelist():
                styles_xml = docx_zip.read('word/styles.xml')
                root = ET.fromstring(styles_xml)
                
                print("\nStili definiti nel documento (da styles.xml):")
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                
                for style in root.findall('.//w:style', ns):
                    style_id = style.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId')
                    style_type = style.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
                    
                    name_elem = style.find('.//w:name', ns)
                    if name_elem is not None:
                        style_name = name_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                        print(f"\n  Stile: {style_name} (ID: {style_id}, Tipo: {style_type})")
                        
                        # Cerca informazioni sul font
                        font_elem = style.find('.//w:rFonts', ns)
                        if font_elem is not None:
                            ascii_font = font_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii')
                            if ascii_font:
                                print(f"    - Font: {ascii_font}")
                        
                        # Cerca dimensione font
                        size_elem = style.find('.//w:sz', ns)
                        if size_elem is not None:
                            size_val = size_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                            if size_val:
                                print(f"    - Dimensione: {int(size_val)/2} pt")
                        
                        # Cerca colore
                        color_elem = style.find('.//w:color', ns)
                        if color_elem is not None:
                            color_val = color_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                            if color_val:
                                print(f"    - Colore: #{color_val}")
            
            # Analizza settings.xml per impostazioni generali
            if 'word/settings.xml' in docx_zip.namelist():
                settings_xml = docx_zip.read('word/settings.xml')
                root = ET.fromstring(settings_xml)
                
                print("\n\nImpostazioni del documento (da settings.xml):")
                
                # Cerca font predefinito
                default_fonts = root.find('.//w:defaultFonts', ns)
                if default_fonts is not None:
                    ascii_theme = default_fonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}asciiTheme')
                    if ascii_theme:
                        print(f"  - Font tema predefinito: {ascii_theme}")
                
    except Exception as e:
        print(f"  Errore nell'analisi XML: {str(e)}")

def analyze_detailed(file_path):
    """Analisi dettagliata del documento"""
    
    if not os.path.exists(file_path):
        print(f"Errore: Il file {file_path} non esiste")
        return
    
    try:
        doc = Document(file_path)
        
        print("\n" + "=" * 80)
        print("ANALISI DETTAGLIATA AGGIUNTIVA")
        print("=" * 80)
        
        # Analisi dettagliata dell'intestazione
        print("\nDETTAGLI INTESTAZIONE:")
        print("-" * 40)
        
        for section in doc.sections:
            header = section.header
            if header.tables:
                for table in header.tables:
                    print(f"\nTabella nell'intestazione:")
                    print(f"  - Numero di righe: {len(table.rows)}")
                    print(f"  - Numero di colonne: {len(table.columns)}")
                    
                    # Analizza ogni cella
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            if cell.text.strip():
                                print(f"\n  Cella [{i},{j}]:")
                                print(f"    - Contenuto: '{cell.text}'")
                                
                                # Analizza i paragrafi nella cella
                                for para in cell.paragraphs:
                                    if para.text.strip():
                                        print(f"    - Allineamento paragrafo: {para.alignment}")
                                        
                                        # Analizza i runs
                                        for run in para.runs:
                                            if run.text.strip():
                                                print(f"\n      Run: '{run.text}'")
                                                if run.font.name:
                                                    print(f"        - Font: {run.font.name}")
                                                if run.font.size:
                                                    print(f"        - Dimensione: {run.font.size.pt} pt")
                                                if run.font.bold:
                                                    print(f"        - Grassetto: Sì")
                                                if run.font.italic:
                                                    print(f"        - Corsivo: Sì")
                                                if run.font.color.rgb:
                                                    rgb = run.font.color.rgb
                                                    print(f"        - Colore: RGB({rgb.red}, {rgb.green}, {rgb.blue})")
        
        # Analisi dettagliata del piè di pagina
        print("\n\nDETTAGLI PIÈ DI PAGINA:")
        print("-" * 40)
        
        for section in doc.sections:
            footer = section.footer
            for i, para in enumerate(footer.paragraphs):
                if para.text.strip():
                    print(f"\nParagrafo {i + 1}:")
                    print(f"  - Contenuto completo: '{para.text}'")
                    print(f"  - Allineamento: {para.alignment}")
                    
                    # Dettagli dei runs
                    for j, run in enumerate(para.runs):
                        if run.text.strip():
                            print(f"\n  Run {j + 1}: '{run.text}'")
                            if run.font.name:
                                print(f"    - Font: {run.font.name}")
                            if run.font.size:
                                print(f"    - Dimensione: {run.font.size.pt} pt")
                            if run.font.bold:
                                print(f"    - Grassetto: Sì")
                            if run.font.italic:
                                print(f"    - Corsivo: Sì")
                            if run.font.underline:
                                print(f"    - Sottolineato: Sì")
                            if run.font.color.rgb:
                                rgb = run.font.color.rgb
                                print(f"    - Colore: RGB({rgb.red}, {rgb.green}, {rgb.blue})")
        
        # Analizza contenuto XML
        analyze_xml_content(file_path)
        
        # Informazioni sui font utilizzati
        print("\n\n8. RIEPILOGO FONT UTILIZZATI:")
        print("-" * 40)
        
        fonts_used = set()
        sizes_used = set()
        
        # Raccogli font dall'intestazione
        for section in doc.sections:
            header = section.header
            for para in header.paragraphs:
                for run in para.runs:
                    if run.font.name:
                        fonts_used.add(run.font.name)
                    if run.font.size:
                        sizes_used.add(run.font.size.pt)
            
            # Font dal piè di pagina
            footer = section.footer
            for para in footer.paragraphs:
                for run in para.runs:
                    if run.font.name:
                        fonts_used.add(run.font.name)
                    if run.font.size:
                        sizes_used.add(run.font.size.pt)
        
        # Font dal contenuto principale
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.name:
                    fonts_used.add(run.font.name)
                if run.font.size:
                    sizes_used.add(run.font.size.pt)
        
        print(f"\nFont utilizzati: {', '.join(sorted(fonts_used)) if fonts_used else 'Nessuno specificato'}")
        print(f"Dimensioni utilizzate: {', '.join([f'{s} pt' for s in sorted(sizes_used)]) if sizes_used else 'Nessuna specificata'}")
        
        print("\n" + "=" * 80)
        
    except Exception as e:
        print(f"Errore durante l'analisi: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    file_path = "/opt/progetti/python/carta_intestata/Modello 2025.docx"
    analyze_detailed(file_path)