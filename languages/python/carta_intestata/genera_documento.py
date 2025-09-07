#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
import argparse
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import markdown
from bs4 import BeautifulSoup
import re

def create_header(document):
    section = document.sections[0]
    header = section.header
    
    table = header.add_table(rows=1, cols=2, width=Cm(18))
    table.autofit = False
    
    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]
    
    left_cell._element.get_or_add_tcPr().append(OxmlElement('w:tcW'))
    left_cell._element.tcPr.tcW.set(qn('w:w'), '5000')
    left_cell._element.tcPr.tcW.set(qn('w:type'), 'pct')
    
    right_cell._element.get_or_add_tcPr().append(OxmlElement('w:tcW'))
    right_cell._element.tcPr.tcW.set(qn('w:w'), '5000')
    right_cell._element.tcPr.tcW.set(qn('w:type'), 'pct')
    
    p1 = left_cell.paragraphs[0]
    run1 = p1.add_run("Paolo Gironi")
    run1.font.name = "Montserrat"
    run1.font.size = Pt(20)
    run1.font.bold = True
    
    p2 = left_cell.add_paragraph()
    run2 = p2.add_run("Web Marketing - specialista SEO e SEM")
    run2.font.name = "Montserrat"
    run2.font.size = Pt(10)
    
    p3 = right_cell.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run3 = p3.add_run("Via Zanella, 4 - 36022 S.Giuseppe di Cassola (VI)")
    run3.font.name = "Montserrat"
    run3.font.size = Pt(8)
    run3.font.bold = True
    
    p4 = right_cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run4 = p4.add_run("Tel. 334 1943957 Partita Iva 03137610246")
    run4.font.name = "Montserrat"
    run4.font.size = Pt(8)
    run4.font.bold = True
    
    p5 = right_cell.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run5 = p5.add_run("C.F. GRNPLA68E19D704Y")
    run5.font.name = "Montserrat"
    run5.font.size = Pt(8)
    run5.font.bold = True
    
    for cell in table.rows[0].cells:
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement('w:' + border_name)
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)
        tcPr.append(tcBorders)
    
    # Aggiungi spazio dopo la tabella dell'intestazione
    p_after = header.add_paragraph()
    p_after.paragraph_format.space_after = Pt(24)

def create_footer(document):
    section = document.sections[0]
    footer = section.footer
    
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("paolo@gironi.it | www.gironi.it")
    run.font.name = "Montserrat"
    run.font.size = Pt(9)

def setup_styles(document):
    styles = document.styles
    
    heading_configs = [
        ("Heading 1", 20, "Montserrat SemiBold"),
        ("Heading 2", 16, "Montserrat SemiBold"),
        ("Heading 3", 14, "Montserrat SemiBold"),
        ("Heading 4", 14, "Proxima Nova"),
        ("Heading 5", 11, "Proxima Nova"),
        ("Heading 6", 11, "Proxima Nova")
    ]
    
    for style_name, size, font in heading_configs:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = font
            style.font.size = Pt(size)
            style.font.bold = True
            
            if style_name in ["Heading 5", "Heading 6"]:
                style.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    if "Normal" in styles:
        normal_style = styles["Normal"]
        normal_style.font.name = "Proxima Nova"
        normal_style.font.size = Pt(11)
        normal_style.paragraph_format.space_after = Pt(12)

def convert_markdown_to_docx(markdown_text, document):
    html = markdown.markdown(markdown_text, extensions=['extra', 'codehilite'])
    soup = BeautifulSoup(html, 'html.parser')
    
    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'blockquote', 'pre']):
        if element.name == 'h1':
            p = document.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            p = document.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            p = document.add_heading(element.get_text(), level=3)
        elif element.name == 'h4':
            p = document.add_heading(element.get_text(), level=4)
        elif element.name == 'h5':
            p = document.add_heading(element.get_text(), level=5)
        elif element.name == 'h6':
            p = document.add_heading(element.get_text(), level=6)
        elif element.name == 'p':
            p = document.add_paragraph()
            process_inline_elements(element, p)
        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li', recursive=False):
                p = document.add_paragraph(style='List Bullet' if element.name == 'ul' else 'List Number')
                process_inline_elements(li, p)
        elif element.name == 'blockquote':
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            process_inline_elements(element, p)
        elif element.name == 'pre':
            code_text = element.get_text()
            p = document.add_paragraph()
            p.add_run(code_text).font.name = "Courier New"

def process_inline_elements(element, paragraph):
    if element.string:
        paragraph.add_run(element.string)
    else:
        for child in element.children:
            if hasattr(child, 'name'):
                if child.name == 'strong' or child.name == 'b':
                    run = paragraph.add_run(child.get_text())
                    run.bold = True
                elif child.name == 'em' or child.name == 'i':
                    run = paragraph.add_run(child.get_text())
                    run.italic = True
                elif child.name == 'code':
                    run = paragraph.add_run(child.get_text())
                    run.font.name = "Courier New"
                else:
                    process_inline_elements(child, paragraph)
            else:
                paragraph.add_run(str(child))

def main():
    parser = argparse.ArgumentParser(description='Genera un documento DOCX con intestazione e piè di pagina da un file Markdown')
    parser.add_argument('input_file', help='File Markdown di input')
    parser.add_argument('-o', '--output', default='output.docx', help='Nome del file DOCX di output (default: output.docx)')
    
    args = parser.parse_args()
    
    try:
        with open(args.input_file, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
    except FileNotFoundError:
        print(f"Errore: Il file '{args.input_file}' non è stato trovato.")
        sys.exit(1)
    except Exception as e:
        print(f"Errore nella lettura del file: {e}")
        sys.exit(1)
    
    document = Document()
    
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(3.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(1)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)
    
    create_header(document)
    create_footer(document)
    setup_styles(document)
    
    convert_markdown_to_docx(markdown_content, document)
    
    try:
        document.save(args.output)
        print(f"Documento generato con successo: {args.output}")
    except Exception as e:
        print(f"Errore nel salvataggio del documento: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()